"""
PCD投稿用 DOCX修正スクリプト
対象: Not_All_Social_Determinants_Travel_20260627.docx
修正内容:
  1. タイトル修正（別論文テンプレートのタイトルが残存）
  2. Abstract: "Background:" → "Aims:" + 200語以内に短縮
  3. Discussion: (b)代替手法 / (d)臨床実践示唆 を追加
  4. 末尾に Funding / Ethics Statement / Declarations of Interest を追加
  5. Data Availability: GitHub/Zenodo参照を追記
"""

import sys
import io

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document

DOCX_PATH = "04_Manuscripts/Not_All_Social_Determinants_Travel_20260627.docx"

doc = Document(DOCX_PATH)


# ─────────────────────────────────────────────
# Helper: 新しい段落XMLエレメントを生成
# ─────────────────────────────────────────────
def make_para_elem(doc_ref, text, style_name, bold_prefix=None):
    """段落を一時的にdocに追加してXMLエレメントを取り出し、本体から切り離して返す。"""
    try:
        p = doc_ref.add_paragraph(style=style_name)
    except KeyError:
        p = doc_ref.add_paragraph()

    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)

    elem = p._element
    elem.getparent().remove(elem)
    return elem


# ═══════════════════════════════════════════════
# 1. タイトル修正 (Para 0)
# ═══════════════════════════════════════════════
new_title = (
    "Not All Social Determinants Travel: "
    "A Prefectural Ecological Analysis of Unemployment Rate "
    "and Diabetes Prevalence in Japan"
)
doc.paragraphs[0].runs[0].text = new_title
print(f"[1] Title → {new_title[:80]}...")


# ═══════════════════════════════════════════════
# 2. Abstract修正 (Para 6–9)
# ═══════════════════════════════════════════════
# Para 6: run[0]="Background:" → "Aims:", run[2]=body text
doc.paragraphs[6].runs[0].text = "Aims:"
doc.paragraphs[6].runs[2].text = (
    "Unemployment is widely recognized as a social determinant of diabetes, "
    "but evidence from low-unemployment settings is limited. "
    "We examined whether prefectural unemployment rates predict diabetes prevalence in Japan"
)
# run[3] は "." のまま

# Para 7: Methods run[2]
doc.paragraphs[7].runs[2].text = (
    "We conducted a prefecture-level ecological study of all 47 Japanese prefectures. "
    "The primary outcome was the proportion with HbA1c ≥ 6.5% (FY2022, NDB Open Data). "
    "The main exposure was the complete unemployment rate (Statistics Bureau, 2022). "
    "We estimated correlations and six sequential OLS models adjusted for aging rate, "
    "per-capita prefectural income, and population density, "
    "with HC3 and outlier-exclusion sensitivity analyses"
)

# Para 8: Results run[2]
doc.paragraphs[8].runs[2].text = (
    "Regional unemployment showed virtually no association with HbA1c high rate "
    "(Pearson r = −0.016, p = 0.916), "
    "and remained non-significant across all six model specifications (all p ≥ 0.495). "
    "Per-capita prefectural income was strongly and inversely associated "
    "(r = −0.570; bivariate R² = 0.325; p < 0.001). "
    "Similar null findings were observed for antidiabetic drug prescriptions"
)

# Para 9: Conclusions run[2]
doc.paragraphs[9].runs[2].text = (
    "In Japan’s low-unemployment environment, income—not unemployment—"
    "was the meaningful ecological predictor of diabetes prevalence. "
    "These findings caution against regarding social determinants as universally "
    "transportable predictors across different socioeconomic contexts"
)

print("[2] Abstract: Background→Aims, trimmed to ≤200 words")


# ═══════════════════════════════════════════════
# 3. Discussion: 代替手法(b) & 臨床実践(d) を追加
#    "Strengths and Limitations" 見出し (Para 78) の直前に挿入
# ═══════════════════════════════════════════════
alt_methods_text = (
    "Alternative methodologies better suited to testing the unemployment–diabetes "
    "hypothesis would include multilevel longitudinal studies linking individual-level "
    "employment histories to HbA1c trajectories over time, "
    "or natural experiment designs exploiting regional economic shocks as instrumental "
    "variables. Prefecture-level ecological analysis, while appropriate for this "
    "publicly available open-data resource, cannot establish individual-level causal "
    "relationships and is susceptible to the ecological fallacy."
)
clinical_text = (
    "From a primary care perspective, these findings suggest that clinicians in Japan "
    "may benefit from prioritising income-based risk screening over employment-status "
    "screening when identifying patients at high risk for poor glycaemic control. "
    "Incorporating enquiry about financial hardship into existing specific health "
    "checkup programmes may enhance the identification of high-risk individuals and "
    "support targeted preventive interventions in the primary care setting."
)

# "Strengths and Limitations" 見出しを探す
p_strengths = None
for p in doc.paragraphs:
    if p.text.strip() == "Strengths and Limitations":
        p_strengths = p
        break

if p_strengths is None:
    print("WARNING: 'Strengths and Limitations' heading not found!")
else:
    cli_elem = make_para_elem(doc, clinical_text, "Body Text")
    alt_elem = make_para_elem(doc, alt_methods_text, "Body Text")
    p_strengths._element.addprevious(cli_elem)
    cli_elem.addprevious(alt_elem)
    print("[3] Discussion: alternative methodologies + clinical practice paragraphs added")


# ═══════════════════════════════════════════════
# 4. Data Availability テキスト末尾に GitHub/Zenodo 追記
#    Para 86 run[4] のテキストを更新
# ═══════════════════════════════════════════════
da_para = None
for p in doc.paragraphs:
    if p.text.startswith("The NDB Open Data used in this study"):
        da_para = p
        break

if da_para and len(da_para.runs) >= 5:
    da_para.runs[4].text = (
        ". Analysis scripts and the integrated dataset are archived on GitHub "
        "(https://github.com/haruki00430/NDB_XXX_diabetes_unemployment) "
        "and on Zenodo (DOI to be assigned upon repository release)."
    )
    print("[4] Data Availability: GitHub/Zenodo reference added")
else:
    print("WARNING: Data Availability paragraph or expected run not found; skipped.")


# ═══════════════════════════════════════════════
# 5. 末尾に Funding / Ethics Statement / Declarations of Interest を追加
#    既存の "Declaration of Generative AI..." 見出しの直前に挿入
# ═══════════════════════════════════════════════
ai_heading = None
for p in doc.paragraphs:
    if p.text.startswith("Declaration of Generative AI"):
        ai_heading = p
        break

if ai_heading is None:
    print("WARNING: 'Declaration of Generative AI' heading not found!")
else:
    H1 = "Heading 1"
    FP = "First Paragraph"

    # --- Declarations of Interest ---
    coi_body = make_para_elem(doc, "The authors declare no competing interests.", FP)
    coi_head = make_para_elem(doc, "Declarations of Interest", H1)
    ai_heading._element.addprevious(coi_body)
    coi_body.addprevious(coi_head)

    # --- Ethics Statement ---
    ethics_body_text = (
        "This study analyzed only publicly available aggregate statistical data "
        "from Japan's National Database of Health Insurance Claims Open Data "
        "(NDB Open Data) and did not involve individual patient data or direct "
        "human subject research. "
        "Ethical approval was not required under the Ethical Guidelines for "
        "Medical and Biological Research Involving Human Subjects (Ministry of "
        "Education, Culture, Sports, Science and Technology; Ministry of Health, "
        "Labour and Welfare; Ministry of Economy, Trade and Industry, Japan, 2021). "
        "Informed consent was not applicable."
    )
    ethics_body = make_para_elem(doc, ethics_body_text, FP)
    ethics_head = make_para_elem(doc, "Ethics Statement", H1)
    coi_head.addprevious(ethics_body)
    ethics_body.addprevious(ethics_head)

    # --- Funding ---
    funding_body_text = (
        "This research did not receive any specific grant from funding agencies "
        "in the public, commercial, or not-for-profit sectors."
    )
    funding_body = make_para_elem(doc, funding_body_text, FP)
    funding_head = make_para_elem(doc, "Funding", H1)
    ethics_head.addprevious(funding_body)
    funding_body.addprevious(funding_head)

    print("[5] Funding / Ethics Statement / Declarations of Interest added")


# ═══════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════
doc.save(DOCX_PATH)
print(f"\n✅ Saved: {DOCX_PATH}")

# 検証: 段落リストを簡易表示
print("\n--- 最終段落構造（抜粋: 先頭 / Abstract / 末尾）---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    if i < 12 or (78 <= i <= 95) or i > len(doc.paragraphs) - 25:
        print(f"  [{i:03d}] {p.style.name:25s} | {t[:90]}")
