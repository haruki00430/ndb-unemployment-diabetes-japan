"""
Phase 4: DOCXに図を挿入
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

project_root = Path(__file__).resolve().parents[3]
PROJECT = "NDB_XXX_diabetes_unemployment"

doc_path = project_root / f"projects/{PROJECT}/04_Manuscripts/Manuscript_diabetes_unemployment.docx"
fig_dir = project_root / f"projects/{PROJECT}/results/figures"

FIGURES = [
    ("Fig1_unemployment_hba1c_scatter.png",
     "Figure 1. Scatter plot of unemployment rate and HbA1c high rate (≥6.5%) across 47 Japanese prefectures. The regression line (solid) with 95% confidence interval (shaded) shows no association (r = −0.016, p = 0.916). Okinawa (highest unemployment rate) is highlighted in red."),
    ("Fig2_unemployment_drug_scatter.png",
     "Figure 2. Scatter plot of unemployment rate and antidiabetic drug prescription quantity (per 10,000 population) across 47 Japanese prefectures (r = −0.103, p = 0.490)."),
    ("Fig3_forest_sensitivity_hba1c.png",
     "Figure 3. Forest plot of sensitivity analysis: unemployment rate regression coefficients (β) with 95% confidence intervals across six model specifications. All coefficients span zero, confirming a consistent null result."),
    ("Fig4_prefecture_ranking.png",
     "Figure 4. Prefectural rankings of unemployment rate (left) and HbA1c high rate (right). The absence of a consistent pattern between the two rankings is consistent with the null ecological association."),
]

doc = Document(str(doc_path))

# References セクションの前に図を追加
ref_para_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() in ["References", "# References"]:
        ref_para_idx = i
        break

if ref_para_idx is None:
    ref_para_idx = len(doc.paragraphs)

insert_idx = ref_para_idx

for fig_file, caption in FIGURES:
    fig_path = fig_dir / fig_file
    if not fig_path.exists():
        print(f"Figure not found: {fig_path}")
        continue

    # 段落を挿入
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(fig_path), width=Inches(5.5))

    cap_para = doc.add_paragraph()
    cap_run = cap_para.add_run(caption)
    cap_run.font.size = Pt(9)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacing

    print(f"Inserted: {fig_file}")

out_path = project_root / f"projects/{PROJECT}/04_Manuscripts/Manuscript_diabetes_unemployment_with_figs.docx"
doc.save(str(out_path))
print(f"\nSaved: {out_path}")
